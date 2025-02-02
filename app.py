import streamlit as st
import requests
from docx import Document
import os

# Configura la API key de Together AI
api_key = st.secrets["together_api_key"]
url = "https://api.together.xyz/v1/chat/completions"
headers = {
    "Authorization": f"Bearer {api_key}",
    "Content-Type": "application/json"
}

# Función para generar la trama y los personajes
def generar_trama_y_personajes(genero, titulo):
    prompt = f"Genera una trama y describe los personajes principales para una novela de género {genero} titulada '{titulo}'. Incluye una tabla de contenidos con 24 capítulos."
    data = {
        "model": "meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
        "messages": [
            {"role": "system", "content": "Eres un escritor experto en novelas."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1,
        "stop": ["<|eot_id|>", "<|eom_id|>"],
        "stream": False  # Desactivamos el streaming para obtener la respuesta completa
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        st.error(f"Error al generar la trama: {response.status_code} - {response.text}")
        return None

# Función para generar un capítulo
def generar_capitulo(trama, capitulo_numero):
    prompt = f"Escribe el capítulo {capitulo_numero} de la novela. La trama general es: {trama}. El capítulo debe tener alrededor de 2000 palabras y debe incluir desarrollo de personajes, descripciones detalladas, subtramas, diálogos extensos, reflexiones internas, eventos detallados, flashbacks y expansión del mundo."
    data = {
        "model": "meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
        "messages": [
            {"role": "system", "content": "Eres un escritor experto en novelas."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1,
        "stop": ["<|eot_id|>", "<|eom_id|>"],
        "stream": False  # Desactivamos el streaming para obtener la respuesta completa
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        st.error(f"Error al generar el capítulo {capitulo_numero}: {response.status_code} - {response.text}")
        return None

# Función para guardar la novela en un archivo .docx
def guardar_novela(titulo, capitulos):
    doc = Document()
    doc.add_heading(titulo, 0)
    for capitulo in capitulos:
        doc.add_heading(f"Capítulo {capitulo['numero']}", level=1)
        doc.add_paragraph(capitulo['contenido'])
    doc.save(f"{titulo}.docx")

# Interfaz de Streamlit
st.title("Generador de Novelas")

# Entradas del usuario
genero = st.text_input("Especifica el género de la novela:")
titulo = st.text_input("Especifica el título de la novela:")

if st.button("Generar Novela"):
    if genero and titulo:
        with st.spinner("Generando trama y personajes..."):
            trama_y_personajes = generar_trama_y_personajes(genero, titulo)
            if trama_y_personajes:
                st.write("### Trama y Personajes")
                st.write(trama_y_personajes)

                capitulos = []
                for i in range(1, 25):
                    with st.spinner(f"Generando capítulo {i}..."):
                        capitulo = generar_capitulo(trama_y_personajes, i)
                        if capitulo:
                            capitulos.append({"numero": i, "contenido": capitulo})
                            st.write(f"### Capítulo {i}")
                            st.write(capitulo)

                # Guardar la novela en un archivo .docx
                guardar_novela(titulo, capitulos)
                st.success("¡Novela generada con éxito!")
                st.markdown(f"### [Descargar novela completa en .docx]({titulo}.docx)")
    else:
        st.error("Por favor, introduce el género y el título de la novela.")
